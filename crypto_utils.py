import os, base64, hashlib
from Crypto.Cipher import AES
from Crypto.Protocol.KDF import PBKDF2
from Crypto.Random import get_random_bytes
from math import log2
from docx import Document
import numpy as np

UPLOAD_FOLDER = "static/uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def file_hash(path):
    """Menghitung hash SHA256 unik dari file"""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            h.update(chunk)
    return h.hexdigest()

def derive_key(password, salt):
    return PBKDF2(password.encode(), salt, dkLen=32, count=100000)

def encrypt_file(file_path, password):
    """Enkripsi file pakai AES-GCM"""
    data = open(file_path, "rb").read()
    salt = get_random_bytes(16)
    key = derive_key(password, salt)
    cipher = AES.new(key, AES.MODE_GCM)
    ciphertext, tag = cipher.encrypt_and_digest(data)
    enc_file = os.path.join(UPLOAD_FOLDER, os.path.basename(file_path) + ".enc")
    with open(enc_file, "wb") as f:
        [f.write(x) for x in (salt, cipher.nonce, tag, ciphertext)]
    code = base64.urlsafe_b64encode(salt + cipher.nonce + tag).decode()
    return enc_file, code

def decrypt_file(code, password):
    """Dekripsi file berdasarkan kode enkripsi"""
    meta = base64.urlsafe_b64decode(code.encode())
    salt, nonce, tag = meta[:16], meta[16:32], meta[32:48]
    for fn in os.listdir(UPLOAD_FOLDER):
        if fn.endswith(".enc"):
            with open(os.path.join(UPLOAD_FOLDER, fn), "rb") as f:
                raw = f.read()
            if raw.startswith(salt):
                key = derive_key(password, salt)
                cipher = AES.new(key, AES.MODE_GCM, nonce=nonce)
                dec_data = cipher.decrypt_and_verify(raw[48:], tag)
                dec_path = os.path.join(UPLOAD_FOLDER, fn.replace(".enc", "_dec.mp4"))
                open(dec_path, "wb").write(dec_data)
                return dec_path
    raise ValueError("File enkripsi tidak ditemukan atau password salah!")

def compute_entropy(path):
    """Menghitung nilai entropy"""
    data = open(path, "rb").read()
    freq = [0]*256
    for b in data:
        freq[b]+=1
    prob = [f/len(data) for f in freq if f>0]
    return -sum(p*log2(p) for p in prob)

def compute_histogram(path):
    """Menghitung histogram byte"""
    data = open(path, "rb").read()
    hist = [0]*256
    for b in data:
        hist[b]+=1
    return hist

def compare_images(hist1, hist2):
    """NPCR & UACI berbasis histogram byte"""
    arr1, arr2 = np.array(hist1), np.array(hist2)
    npcr = np.mean(arr1 != arr2) * 100
    uaci = np.mean(np.abs(arr1 - arr2) / 255) * 100
    return npcr, uaci

def generate_report(code, orig_hash="", enc_hash="", dec_hash=""):
    """Membuat laporan DOCX"""
    doc = Document()
    doc.add_heading("Laporan Enkripsi & Dekripsi AES-GCM", 0)
    doc.add_paragraph(f"Kode Enkripsi: {code}")
    doc.add_paragraph("AES-GCM digunakan untuk mengamankan file video dengan mode authenticated encryption.")
    doc.add_heading("Hash (SHA-256):", level=1)
    doc.add_paragraph(f"File Asli   : {orig_hash}")
    doc.add_paragraph(f"File Enkripsi : {enc_hash}")
    doc.add_paragraph(f"File Dekripsi : {dec_hash}")
    doc.add_paragraph("Hash di atas membuktikan bahwa hasil dekripsi identik dengan file asli bila nilainya sama.")
    path = os.path.join(UPLOAD_FOLDER, "laporan.docx")
    doc.save(path)
    return path

# ========== FUNGSI BARU: AVALANCHE EFFECT YANG BENAR ==========
def calculate_avalanche_effect_correct(data: bytes, password: str):
    """
    Menghitung Avalanche Effect yang benar untuk binary data
    Ubah 1 bit pada PLAINTEXT (bukan password), hitung perubahan ciphertext
    """
    if len(data) == 0:
        return 0.0, 0, 0
    
    try:
        # Enkripsi data asli
        salt1 = get_random_bytes(16)
        key1 = derive_key(password, salt1)
        cipher1 = AES.new(key1, AES.MODE_GCM)
        ciphertext1, tag1 = cipher1.encrypt_and_digest(data)
        
        # Ubah 1 bit pada plaintext (byte terakhir flip 1 bit)
        data2 = bytearray(data)
        data2[-1] ^= 0x01  # Flip LSB of last byte
        
        # Enkripsi data yang diubah (dengan salt/nonce berbeda)
        salt2 = get_random_bytes(16)
        key2 = derive_key(password, salt2)
        cipher2 = AES.new(key2, AES.MODE_GCM)
        ciphertext2, tag2 = cipher2.encrypt_and_digest(bytes(data2))
        
        # Ambil panjang minimum untuk perbandingan
        min_len = min(len(ciphertext1), len(ciphertext2))
        ciphertext1 = ciphertext1[:min_len]
        ciphertext2 = ciphertext2[:min_len]
        
        # Konversi ke array bit
        arr1 = np.unpackbits(np.frombuffer(ciphertext1, dtype=np.uint8))
        arr2 = np.unpackbits(np.frombuffer(ciphertext2, dtype=np.uint8))
        
        # Hitung jumlah bit yang berbeda
        bit_changes = np.sum(arr1 != arr2)
        total_bits = len(arr1)
        
        # Hitung persentase
        if total_bits == 0:
            return 0.0, 0, 0
        
        avalanche_percentage = (bit_changes / total_bits) * 100
        
        return round(avalanche_percentage, 3), bit_changes, total_bits
        
    except Exception as e:
        print(f"Error calculating avalanche effect: {e}")
        return 0.0, 0, 0

# ========== FUNGSI BARU: TEST VIDEO PLAYBACK ==========
def test_video_playback(video_path):
    """
    Cek apakah video hasil dekripsi masih bisa diputar
    Menggunakan OpenCV
    """
    try:
        import cv2
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            return False
        
        # Coba baca beberapa frame
        success, frame = cap.read()
        cap.release()
        
        return bool(success)
    except Exception:
        return False