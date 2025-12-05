from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session
import os, time, base64, hashlib, secrets, glob
from Crypto.Cipher import AES
from Crypto.Protocol.KDF import PBKDF2
import numpy as np
from io import BytesIO
from docx import Document
from pathlib import Path
from werkzeug.utils import secure_filename
import cv2  # Tambahan untuk video processing

app = Flask(__name__)
app.secret_key = 'aes-demo-secret'
app.config['SESSION_TYPE'] = 'filesystem'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # Batasi upload 50MB
app.config['PERMANENT_SESSION_LIFETIME'] = 1800  # 30 menit

UPLOAD_FOLDER = 'static/uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ========== CUSTOM TEMPLATE FILTER ==========
@app.template_filter('filetype')
def filetype_filter(filename):
    """Custom filter untuk menentukan tipe file di template"""
    if not filename:
        return "File"
    
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    if ext in ['mp4', 'avi', 'mov', 'mkv', 'webm', 'flv', 'wmv']:
        return "Video"
    elif ext in ['png', 'jpg', 'jpeg', 'bmp', 'gif', 'webp']:
        return "Gambar"
    else:
        return "File"

# ========== FUNGSI CLEANUP ==========
def cleanup_temp_files():
    """Hapus file temporary .orig saat aplikasi berhenti"""
    pattern = os.path.join(UPLOAD_FOLDER, "*.orig")
    for file_path in glob.glob(pattern):
        try:
            os.remove(file_path)
            print(f"Cleaned up: {file_path}")
        except:
            pass

import atexit
atexit.register(cleanup_temp_files)

# ========== Fungsi utilitas ==========
def sha256sum(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()

def entropy(data: bytes) -> float:
    if not data:
        return 0.0
    p, _ = np.histogram(np.frombuffer(data, dtype=np.uint8), bins=256, range=(0, 256))
    p = p / np.sum(p)
    p = p[p > 0]
    return -np.sum(p * np.log2(p))

def histogram_data(data: bytes):
    """Hitung histogram dan return sebagai list Python (bukan numpy array)"""
    if not data:
        return [0] * 256
    
    try:
        # Konversi bytes ke numpy array
        arr = np.frombuffer(data, dtype=np.uint8)
        # Hitung histogram
        hist, _ = np.histogram(arr, bins=256, range=(0, 256))
        # Konversi ke list Python (penting untuk JSON serialization)
        return hist.tolist()
    except Exception as e:
        print(f"Error in histogram_data: {e}")
        return [0] * 256

# ========== FUNGSI KORELASI TANPA MINUS ==========
def encryption_correlation(data1, data2, return_details=False):
    """
    Menghitung korelasi untuk analisis kriptografi
    Menggunakan nilai ABSOLUT (tanpa minus) karena untuk enkripsi:
    - Korelasi -0.8 sama buruknya dengan +0.8
    - Yang diukur adalah BESARNYA hubungan, bukan arahnya
    
    Jika return_details=True: return (magnitude, quality, raw_value)
    """
    if len(data1) != len(data2):
        min_len = min(len(data1), len(data2))
        data1 = data1[:min_len]
        data2 = data2[:min_len]
    
    if len(data1) < 2:  # Minimal 2 data untuk korelasi
        if return_details:
            return 0.0, "Insufficient data", 0.0
        return 0.0
    
    try:
        arr1 = np.frombuffer(data1, dtype=np.uint8).astype(float)
        arr2 = np.frombuffer(data2, dtype=np.uint8).astype(float)
        
        # Hitung korelasi Pearson standar
        correlation_matrix = np.corrcoef(arr1, arr2)
        raw_correlation = correlation_matrix[0, 1]
        
        # Handle NaN atau data konstan
        if np.isnan(raw_correlation):
            if return_details:
                return 0.0, "Constant data (no variation)", 0.0
            return 0.0
        
        # Untuk analisis kriptografi: gunakan NILAI ABSOLUT
        # Karena korelasi -0.9 sama buruknya dengan +0.9
        correlation_magnitude = abs(raw_correlation)
        
        # Interpretasi kualitas
        if correlation_magnitude < 0.01:
            quality = "Excellent (practically uncorrelated)"
            quality_short = "Sangat Baik"
        elif correlation_magnitude < 0.05:
            quality = "Very Good (almost no correlation)"
            quality_short = "Baik"
        elif correlation_magnitude < 0.1:
            quality = "Good (very weak correlation)"
            quality_short = "Cukup Baik"
        elif correlation_magnitude < 0.2:
            quality = "Fair (weak correlation)"
            quality_short = "Cukup"
        elif correlation_magnitude < 0.3:
            quality = "Poor (moderate correlation)"
            quality_short = "Buruk"
        else:
            quality = "Very Poor (strong correlation)"
            quality_short = "Sangat Buruk"
        
        if return_details:
            return (
                round(correlation_magnitude, 6),
                quality_short,
                round(raw_correlation, 6),
                quality,
                "negative" if raw_correlation < 0 else "positive"
            )
        else:
            return round(correlation_magnitude, 6)
            
    except Exception as e:
        print(f"Error in encryption_correlation: {e}")
        if return_details:
            return 0.0, "Error in calculation", 0.0
        return 0.0

# ========== FUNGSI AVALANCHE EFFECT YANG DIPERBAIKI ==========
def calculate_avalanche_effect(data: bytes, password: str):
    """
    Menghitung Avalanche Effect yang BENAR untuk binary data
    Ubah 1 bit pada PLAINTEXT (bukan password), hitung % perubahan ciphertext
    """
    if len(data) == 0:
        return 0.0, 0, 0
    
    try:
        # Enkripsi data asli
        enc_data1, _ = aes_encrypt(data, password)
        
        # Ubah 1 bit pada plaintext (byte terakhir flip LSB)
        data2 = bytearray(data)
        data2[-1] ^= 0x01  # Flip LSB of last byte
        
        # Enkripsi data yang diubah
        enc_data2, _ = aes_encrypt(bytes(data2), password)
        
        # Pastikan panjang sama (ambil minimum)
        min_len = min(len(enc_data1), len(enc_data2))
        enc_data1 = enc_data1[:min_len]
        enc_data2 = enc_data2[:min_len]
        
        # Konversi ke array bit
        arr1 = np.unpackbits(np.frombuffer(enc_data1, dtype=np.uint8))
        arr2 = np.unpackbits(np.frombuffer(enc_data2, dtype=np.uint8))
        
        # Hitung jumlah bit yang berbeda
        bit_changes = np.sum(arr1 != arr2)
        total_bits = len(arr1)
        
        # Hitung persentase (ideal ~50%)
        if total_bits == 0:
            return 0.0, 0, 0
        
        avalanche_percentage = (bit_changes / total_bits) * 100
        
        return round(avalanche_percentage, 3), bit_changes, total_bits
        
    except Exception as e:
        print(f"Error calculating AE: {e}")
        return 0.0, 0, 0

# ========== FUNGSI VIDEO METRICS ==========
def calculate_video_metrics(video_path):
    """
    Menghitung Frame Processing Time dan FPS untuk video
    Hanya untuk file video yang valid (bukan file terenkripsi)
    """
    if not is_video(video_path):
        return None
    
    try:
        cap = cv2.VideoCapture(video_path)
        
        # Dapatkan informasi video
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps_original = cap.get(cv2.CAP_PROP_FPS)
        
        if total_frames == 0 or fps_original == 0:
            cap.release()
            return None
        
        # Hitung waktu processing per frame (simulasi)
        start_time = time.time()
        frame_count = 0
        max_frames_to_process = min(30, total_frames)  # Process max 30 frames untuk speed
        
        while frame_count < max_frames_to_process:
            ret, frame = cap.read()
            if not ret:
                break
            frame_count += 1
        
        elapsed_time = time.time() - start_time
        cap.release()
        
        if frame_count == 0:
            return None
        
        # Hitung metrics
        avg_frame_time = (elapsed_time / frame_count) * 1000  # dalam ms
        processing_fps = frame_count / elapsed_time if elapsed_time > 0 else 0
        duration = total_frames / fps_original if fps_original > 0 else 0
        
        return {
            'frame_time': round(avg_frame_time, 2),  # ms per frame
            'fps': round(processing_fps, 2),  # frames per second saat processing
            'total_frames': total_frames,
            'duration': round(duration, 2),  # durasi video dalam detik
            'original_fps': round(fps_original, 2)
        }
    except Exception as e:
        print(f"Error calculating video metrics: {e}")
        return None

# ========== FUNGSI TEST VIDEO PLAYBACK ==========
def test_video_playback(video_path):
    """
    Cek apakah video hasil dekripsi masih bisa diputar
    """
    try:
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            return False
        
        # Coba baca beberapa frame
        success, frame = cap.read()
        cap.release()
        
        return bool(success)
    except Exception as e:
        print(f"Error testing video playback: {e}")
        return False

def is_image(filename):
    ext = filename.lower().split('.')[-1]
    return ext in ['png', 'jpg', 'jpeg', 'bmp', 'gif', 'webp']

def is_video(filename):
    ext = filename.lower().split('.')[-1]
    return ext in ['mp4', 'avi', 'mov', 'mkv', 'webm', 'flv', 'wmv']

# ========== Fungsi AES ==========
def aes_encrypt(data: bytes, password: str):
    salt = os.urandom(16)
    key = PBKDF2(password, salt, dkLen=32, count=100000)
    cipher = AES.new(key, AES.MODE_GCM)
    ciphertext, tag = cipher.encrypt_and_digest(data)
    result = salt + cipher.nonce + tag + ciphertext
    code = base64.urlsafe_b64encode(salt + cipher.nonce + tag).decode()
    return result, code

def aes_decrypt(data: bytes, password: str):
    salt, nonce, tag = data[:16], data[16:32], data[32:48]
    ciphertext = data[48:]
    key = PBKDF2(password, salt, dkLen=32, count=100000)
    cipher = AES.new(key, AES.MODE_GCM, nonce=nonce)
    plaintext = cipher.decrypt_and_verify(ciphertext, tag)
    return plaintext

# ========== ROUTES ==========
@app.route('/')
def index():
    return render_template('index.html',
                         hist_orig=[],
                         hist_enc=[], 
                         hist_dec=[],
                         orig_hash='-',
                         enc_hash='-',
                         dec_hash='-',
                         orig_size=0,
                         enc_size=0,
                         dec_size=0,
                         ent_orig='-',
                         ent_enc='-',
                         ent_dec='-',
                         enc_time=0,
                         dec_time=0,
                         code='-',
                         preview_orig=None,
                         preview_dec=None,
                         orig_time=0,
                         correlation_enc='-',
                         correlation_enc_quality='-',
                         correlation_dec='-',
                         correlation_dec_quality='-',
                         frame_time_orig='-',
                         fps_orig='-',
                         frame_time_enc='-',
                         fps_enc='-',
                         frame_time_dec='-',
                         fps_dec='-',
                         avalanche_effect='-',
                         playback_test='-')

@app.route('/encrypt', methods=['POST'])
def encrypt():
    try:
        file = request.files['file']
        password = request.form['password']
        if not file or not password:
            flash("File atau password tidak boleh kosong!")
            return redirect(url_for('index'))

        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        # Baca data dengan context manager
        with open(filepath, 'rb') as f:
            data = f.read()

        # --- Statistik Asli ---
        orig_hash = sha256sum(data)
        ent_orig = round(entropy(data), 4)
        hist_orig = histogram_data(data)  # Sudah dalam bentuk list Python
        orig_size = len(data)

        # --- VIDEO METRICS ASLI ---
        video_metrics_orig = None
        if is_video(filename):
            video_metrics_orig = calculate_video_metrics(filepath)
        
        frame_time_orig = video_metrics_orig['frame_time'] if video_metrics_orig else '-'
        fps_orig = video_metrics_orig['fps'] if video_metrics_orig else '-'

        # --- Enkripsi ---
        start = time.time()
        enc_data, code = aes_encrypt(data, password)
        enc_time = time.time() - start

        # Simpan file terenkripsi
        enc_path = os.path.join(UPLOAD_FOLDER, f"{filename}.enc")
        with open(enc_path, 'wb') as f:
            f.write(enc_data)
        
        # Simpan backup file asli untuk dekripsi nanti
        orig_backup_path = os.path.join(UPLOAD_FOLDER, f"{filename}.orig")
        with open(orig_backup_path, 'wb') as f:
            f.write(data)

        enc_hash = sha256sum(enc_data)
        ent_enc = round(entropy(enc_data), 4)
        hist_enc = histogram_data(enc_data)  # Sudah dalam bentuk list Python
        enc_size = len(enc_data)

        # --- HITUNG KORELASI ENKRIPSI (TANPA MINUS) ---
        correlation_enc, correlation_enc_quality, raw_corr_enc, _, corr_direction = encryption_correlation(data, enc_data, return_details=True)

        # --- HITUNG AVALANCHE EFFECT YANG BENAR ---
        avalanche_effect, bit_changes, total_bits = calculate_avalanche_effect(data, password)

        # Preview
        preview_orig = None
        if is_image(filename) or is_video(filename):
            preview_orig = filename

        # ========== SIMPAN KE SESSION (DENGAN HISTOGRAM) ==========
        session['last_encryption'] = {
            'orig_hash': orig_hash,
            'orig_size': orig_size,
            'ent_orig': ent_orig,
            'enc_hash': enc_hash,
            'enc_size': enc_size,
            'ent_enc': ent_enc,
            'enc_time': enc_time,
            'preview_orig': preview_orig,
            'code': code,
            'orig_path': orig_backup_path,
            'enc_path': enc_path,
            'orig_filename': filename,
            'correlation_enc': correlation_enc,
            'correlation_enc_quality': correlation_enc_quality,
            'raw_correlation_enc': raw_corr_enc,
            'correlation_direction': corr_direction,
            'frame_time_orig': frame_time_orig,
            'fps_orig': fps_orig,
            'avalanche_effect': avalanche_effect,
            'avalanche_bits': f"{bit_changes}/{total_bits}",
            'timestamp': time.time(),
            # SIMPAN HISTOGRAM KE SESSION (SUDAH DALAM BENTUK LIST)
            'hist_orig': hist_orig,
            'hist_enc': hist_enc
        }

        flash("File berhasil dienkripsi!")
        return render_template('index.html',
                               # File Asli
                               orig_hash=orig_hash, 
                               orig_size=orig_size,
                               ent_orig=ent_orig, 
                               hist_orig=hist_orig,
                               preview_orig=preview_orig,
                               frame_time_orig=frame_time_orig,
                               fps_orig=fps_orig,
                               
                               # Enkripsi
                               enc_hash=enc_hash, 
                               enc_size=enc_size,
                               ent_enc=ent_enc, 
                               hist_enc=hist_enc,
                               enc_time=enc_time,
                               code=code,
                               correlation_enc=correlation_enc,
                               correlation_enc_quality=correlation_enc_quality,
                               frame_time_enc='-',
                               fps_enc='-',
                               avalanche_effect=avalanche_effect,
                               
                               # Dekripsi (kosong)
                               dec_hash='-',
                               dec_size=0,
                               ent_dec='-',
                               hist_dec=[],
                               dec_time=0,
                               preview_dec=None,
                               correlation_dec='-',
                               correlation_dec_quality='-',
                               frame_time_dec='-',
                               fps_dec='-',
                               orig_time=0,
                               playback_test='-')

    except Exception as e:
        flash(f"Enkripsi gagal: {str(e)}")
        return redirect(url_for('index'))

@app.route('/decrypt', methods=['POST'])
def decrypt():
    try:
        code = request.form['code']
        password = request.form['password']
        if not code or not password:
            flash("Kode atau password tidak boleh kosong!")
            return redirect(url_for('index'))

        # Ambil data dari SESSION
        last_encryption = session.get('last_encryption')
        if not last_encryption:
            flash("Silakan lakukan enkripsi terlebih dahulu!")
            return redirect(url_for('index'))
        
        # Cek expiry (30 menit)
        timestamp = last_encryption.get('timestamp', 0)
        if time.time() - timestamp > 1800:
            flash("Session expired! Silakan enkripsi ulang.")
            return redirect(url_for('index'))

        # BACA DATA DARI FILE, bukan dari session
        enc_path = last_encryption.get('enc_path')
        if not enc_path or not os.path.exists(enc_path):
            flash("File enkripsi tidak ditemukan!")
            return redirect(url_for('index'))
        
        with open(enc_path, 'rb') as f:
            enc_data = f.read()
            
        orig_path = last_encryption.get('orig_path')
        orig_data = b''
        if orig_path and os.path.exists(orig_path):
            with open(orig_path, 'rb') as f:
                orig_data = f.read()

        start = time.time()
        dec_data = aes_decrypt(enc_data, password)
        dec_time = time.time() - start

        orig_filename = last_encryption.get('orig_filename', 'file')
        if orig_filename.endswith('.enc'):
            dec_filename = orig_filename[:-4] + '_dec'
        else:
            dec_filename = orig_filename + '_dec'
        
        dec_path = os.path.join(UPLOAD_FOLDER, dec_filename)
        with open(dec_path, 'wb') as f:
            f.write(dec_data)

        # Statistik Dekripsi
        dec_hash = sha256sum(dec_data)
        ent_dec = round(entropy(dec_data), 4)
        hist_dec = histogram_data(dec_data)  # Sudah dalam bentuk list Python
        dec_size = len(dec_data)

        # --- HITUNG KORELASI DEKRIPSI (TANPA MINUS) ---
        if orig_data:
            correlation_dec, correlation_dec_quality, _, _, _ = encryption_correlation(orig_data, dec_data, return_details=True)
        else:
            correlation_dec = 0.0
            correlation_dec_quality = "No original data for comparison"

        # --- VIDEO METRICS DEKRIPSI ---
        frame_time_dec = '-'
        fps_dec = '-'
        playback_test = '-'
        
        if is_video(dec_filename):
            video_metrics_dec = calculate_video_metrics(dec_path)
            frame_time_dec = video_metrics_dec['frame_time'] if video_metrics_dec else '-'
            fps_dec = video_metrics_dec['fps'] if video_metrics_dec else '-'
            
            # Test video playback setelah dekripsi
            playback_test = 'BISA' if test_video_playback(dec_path) else 'TIDAK BISA'

        # Preview
        preview_dec = None
        if is_image(dec_filename) or is_video(dec_filename):
            preview_dec = dec_filename

        # AMBIL HISTOGRAM DARI SESSION
        hist_orig_from_session = last_encryption.get('hist_orig', [])
        hist_enc_from_session = last_encryption.get('hist_enc', [])
        
        # Debug print untuk memastikan data histogram ada
        print(f"DEBUG: Histogram dari session - orig: {len(hist_orig_from_session)}, enc: {len(hist_enc_from_session)}")
        print(f"DEBUG: Histogram baru - dec: {len(hist_dec)}")

        flash("File berhasil didekripsi!")
        return render_template('index.html',
                               # Data File Asli (dari session)
                               orig_hash=last_encryption.get('orig_hash', '-'),
                               orig_size=last_encryption.get('orig_size', 0),
                               ent_orig=last_encryption.get('ent_orig', '-'),
                               hist_orig=hist_orig_from_session,  # Ambil dari session
                               preview_orig=last_encryption.get('preview_orig'),
                               frame_time_orig=last_encryption.get('frame_time_orig', '-'),
                               fps_orig=last_encryption.get('fps_orig', '-'),
                               
                               # Data Enkripsi (dari session)
                               enc_hash=last_encryption.get('enc_hash', '-'),
                               enc_size=last_encryption.get('enc_size', 0),
                               ent_enc=last_encryption.get('ent_enc', '-'),
                               hist_enc=hist_enc_from_session,  # Ambil dari session
                               enc_time=last_encryption.get('enc_time', 0),
                               code=last_encryption.get('code', '-'),
                               correlation_enc=last_encryption.get('correlation_enc', '-'),
                               correlation_enc_quality=last_encryption.get('correlation_enc_quality', '-'),
                               frame_time_enc='-',
                               fps_enc='-',
                               avalanche_effect=last_encryption.get('avalanche_effect', '-'),
                               
                               # Data Dekripsi (baru dihitung)
                               dec_hash=dec_hash, 
                               ent_dec=ent_dec, 
                               hist_dec=hist_dec,  # Histogram baru
                               dec_size=dec_size, 
                               dec_time=dec_time,
                               preview_dec=preview_dec,
                               correlation_dec=correlation_dec,
                               correlation_dec_quality=correlation_dec_quality,
                               frame_time_dec=frame_time_dec,
                               fps_dec=fps_dec,
                               orig_time=0,
                               playback_test=playback_test)

    except Exception as e:
        flash(f"Dekripsi gagal: {str(e)}")
        return redirect(url_for('index'))

@app.route('/download/<filename>')
def download_dec(filename):
    return send_file(os.path.join(UPLOAD_FOLDER, filename), as_attachment=True)

@app.route('/report', methods=['POST'])
def report():
    code = request.form.get('code', '')
    doc = Document()
    doc.add_heading('Laporan Enkripsi AES-GCM', 0)
    doc.add_paragraph(f'Kode Unik: {code}')
    doc.add_paragraph('Laporan ini dihasilkan otomatis oleh sistem Flask.')
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return send_file(out, as_attachment=True, download_name='laporan_AES.docx')

@app.route('/cleanup', methods=['POST'])
def cleanup():
    """Route khusus untuk membersihkan session dan file temporary"""
    try:
        # Hapus file temporary
        pattern = os.path.join(UPLOAD_FOLDER, "*.orig")
        for file_path in glob.glob(pattern):
            try:
                os.remove(file_path)
            except:
                pass
        
        # Clear session
        session.clear()
        flash("Session dan file temporary berhasil dibersihkan!")
    except Exception as e:
        flash(f"Cleanup gagal: {str(e)}")
    
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)