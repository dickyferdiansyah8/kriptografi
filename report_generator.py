# report_generator.py
# Generate DOCX report for metrics. Optional PDF conversion if docx2pdf available.
from docx import Document
from docx.shared import Pt
import os

def generate_docx_report(out_path: str, title: str, metrics: dict, hist_samples: dict = None):
    """
    out_path: path to save .docx
    metrics: dict of key->value (MSE, PSNR, NPCR, UACI, entropy, sizes, hashes, times)
    hist_samples: optional dict with 'orig','enc','dec' -> small sample text or JSON
    """
    doc = Document()
    doc.add_heading(title, level=1)

    doc.add_paragraph("Ringkasan metrik empiris:")
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metrik'
    hdr_cells[1].text = 'Nilai'
    for k, v in metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        row_cells[1].text = str(v)

    if hist_samples:
        doc.add_page_break()
        doc.add_heading("Histogram (ringkasan)", level=2)
        for name, sample in hist_samples.items():
            doc.add_heading(name, level=3)
            doc.add_paragraph(str(sample)[:2000])  # avoid extremely long dumps

    doc.save(out_path)
    return out_path

def try_convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return True
    except Exception:
        return False
